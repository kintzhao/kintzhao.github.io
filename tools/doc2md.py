#!/usr/bin/env python3
"""
DOCX to Markdown 转换工具

将指定目录下的所有 .docx 文件转换为 .md 文件

依赖安装：
    pip install python-docx

使用方法：
    python doc2md.py [目录路径]
    
示例：
    python doc2md.py ../content/llm/模型压缩与知识蒸馏
    python doc2md.py ../content/llm  # 递归处理所有子目录
"""

import os
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误：缺少 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


def convert_docx_to_md(docx_path: Path, output_path: Path = None) -> str:
    """
    将 .docx 文件转换为 Markdown 格式
    
    Args:
        docx_path: docx 文件路径
        output_path: 输出 md 文件路径（默认同名 .md）
    
    Returns:
        生成的 Markdown 文本
    """
    if output_path is None:
        output_path = docx_path.with_suffix('.md')
    
    doc = Document(str(docx_path))
    md_lines = []
    
    # 从文件名提取标题
    filename = docx_path.stem
    title = filename.replace('_', ' ').replace('-', ' ')
    
    # 添加 frontmatter
    md_lines.append('---')
    md_lines.append(f'title: {title}')
    md_lines.append(f'source: {docx_path.name}')
    md_lines.append('---')
    md_lines.append('')
    
    # 处理文档内容
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append('')
            continue
        
        # 判断标题级别
        style_name = para.style.name.lower() if para.style else ''
        
        if 'heading 1' in style_name or 'title' in style_name:
            md_lines.append(f'# {text}')
        elif 'heading 2' in style_name:
            md_lines.append(f'## {text}')
        elif 'heading 3' in style_name:
            md_lines.append(f'### {text}')
        elif 'heading 4' in style_name:
            md_lines.append(f'#### {text}')
        elif 'heading 5' in style_name:
            md_lines.append(f'##### {text}')
        elif 'heading 6' in style_name:
            md_lines.append(f'###### {text}')
        else:
            # 检查是否是标题格式（纯文本中常见的格式）
            if text.startswith('第') and ('章' in text or '节' in text):
                md_lines.append(f'## {text}')
            elif re.match(r'^[一二三四五六七八九十]+[、.．]', text):
                md_lines.append(f'## {text}')
            elif re.match(r'^\d+[、.．]\s*\S', text):
                md_lines.append(f'### {text}')
            elif re.match(r'^\([一二三四五六七八九十]+\)', text):
                md_lines.append(f'#### {text}')
            elif re.match(r'^\(\d+\)', text):
                md_lines.append(f'#### {text}')
            else:
                # 普通段落
                md_lines.append(text)
    
    # 处理表格
    for table in doc.tables:
        md_lines.append('')
        table_md = convert_table_to_md(table)
        md_lines.extend(table_md)
        md_lines.append('')
    
    md_content = '\n'.join(md_lines)
    
    # 清理多余空行
    md_content = re.sub(r'\n{3,}', '\n\n', md_content)
    
    # 写入文件
    output_path.write_text(md_content, encoding='utf-8')
    print(f"转换完成: {docx_path.name} -> {output_path.name}")
    
    return md_content


def convert_table_to_md(table) -> list:
    """
    将 Word 表格转换为 Markdown 表格
    """
    if not table.rows:
        return []
    
    md_lines = []
    
    # 获取所有行数据
    rows_data = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows_data.append(cells)
    
    if not rows_data:
        return []
    
    # 计算每列最大宽度
    num_cols = max(len(row) for row in rows_data)
    col_widths = []
    for i in range(num_cols):
        max_width = max(len(row[i]) if i < len(row) else 0 for row in rows_data)
        col_widths.append(max(max_width, 3))
    
    # 生成表格
    for i, row in enumerate(rows_data):
        # 填充缺失的列
        while len(row) < num_cols:
            row.append('')
        
        # 生成表格行
        cells = [row[j].ljust(col_widths[j]) for j in range(num_cols)]
        md_lines.append('| ' + ' | '.join(cells) + ' |')
        
        # 第一行后添加分隔线
        if i == 0:
            separator = '| ' + ' | '.join('-' * w for w in col_widths) + ' |'
            md_lines.append(separator)
    
    return md_lines


def process_directory(directory: Path, recursive: bool = True):
    """
    处理目录下的所有 .docx 文件
    
    Args:
        directory: 目录路径
        recursive: 是否递归处理子目录
    """
    if recursive:
        docx_files = list(directory.rglob('*.docx'))
    else:
        docx_files = list(directory.glob('*.docx'))
    
    # 排除临时文件
    docx_files = [f for f in docx_files if not f.name.startswith('~$')]
    
    if not docx_files:
        print(f"未找到 .docx 文件: {directory}")
        return
    
    print(f"找到 {len(docx_files)} 个 .docx 文件")
    
    success_count = 0
    for docx_file in sorted(docx_files):
        try:
            convert_docx_to_md(docx_file)
            success_count += 1
        except Exception as e:
            print(f"转换失败: {docx_file.name} - {e}")
    
    print(f"\n完成！成功转换 {success_count}/{len(docx_files)} 个文件")


def main():
    if len(sys.argv) < 2:
        # 默认处理 content/llm 目录
        script_dir = Path(__file__).parent
        default_dir = script_dir.parent / 'content' / 'llm'
        
        if default_dir.exists():
            print(f"使用默认目录: {default_dir}")
            process_directory(default_dir)
        else:
            print(__doc__)
            print("\n请指定要处理的目录路径")
            sys.exit(1)
    else:
        target_dir = Path(sys.argv[1])
        
        if not target_dir.exists():
            print(f"目录不存在: {target_dir}")
            sys.exit(1)
        
        process_directory(target_dir)


if __name__ == '__main__':
    main()
